"""
Document parser service for extracting text from PDF and Word documents
"""

import io
from typing import Union
import PyPDF2
from docx import Document


class DocumentParser:
    """Service for parsing various document formats"""
    
    async def parse_document(
        self,
        file_content: bytes,
        filename: str
    ) -> str:
        """
        Parse document and extract complete text content
        
        Args:
            file_content: File content as bytes
            filename: Original filename
        
        Returns:
            Complete extracted text content (no truncation)
        """
        try:
            if filename.lower().endswith('.pdf'):
                return await self._parse_pdf(file_content)
            elif filename.lower().endswith(('.docx', '.doc')):
                return await self._parse_word(file_content)
            else:
                raise ValueError(f"Unsupported file format: {filename}")
        
        except Exception as e:
            raise Exception(f"Failed to parse document {filename}: {str(e)}")
    
    async def _parse_pdf(self, file_content: bytes) -> str:
        """
        Parse PDF document and extract all text
        
        Args:
            file_content: PDF file content as bytes
        
        Returns:
            Complete extracted text from all pages
        """
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            
            # Extract text from all pages
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text:
                    # Add page separator for better context
                    text_content.append(f"--- Page {page_num} ---\n{text}")
            
            # Return complete text without any truncation
            full_text = "\n\n".join(text_content)
            return full_text
        
        except Exception as e:
            raise Exception(f"Failed to parse PDF: {str(e)}")
    
    async def _parse_word(self, file_content: bytes) -> str:
        """
        Parse Word document and extract all text
        
        Args:
            file_content: Word file content as bytes
        
        Returns:
            Complete extracted text including tables
        """
        try:
            doc_file = io.BytesIO(file_content)
            doc = Document(doc_file)
            
            text_content = []
            
            # Extract all paragraph text
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract all text from tables
            for table_num, table in enumerate(doc.tables, 1):
                table_text = [f"\n--- Table {table_num} ---"]
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        table_text.append(row_text)
                text_content.append("\n".join(table_text))
            
            # Return complete text without any truncation
            full_text = "\n\n".join(text_content)
            return full_text
        
        except Exception as e:
            raise Exception(f"Failed to parse Word document: {str(e)}")